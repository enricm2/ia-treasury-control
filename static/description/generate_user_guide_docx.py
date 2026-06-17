#!/usr/bin/env python3
"""
Generate IA Treasury Control User Guide (English section) as a professional Word .docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1B, 0x3A, 0x6B)   # H1 / table header
BLUE      = RGBColor(0x2E, 0x6D, 0xA4)   # H2
DARKGRAY  = RGBColor(0x44, 0x44, 0x44)   # H3
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = RGBColor(0xEB, 0xF5, 0xFB)  # alternating table rows
CODE_BG   = RGBColor(0xF4, 0xF4, 0xF4)
QUOTE_BG  = RGBColor(0xFF, 0xFD, 0xE7)   # light yellow
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)
LINE_COLOR = NAVY

LOGO_PATH = "/opt/odoo18/custom-addons/treasury-control/uniasser_icon_256.png"
# VANTIS_LOGO_PATH = "/path/to/vantis_logo.png"  # TODO: add Vantis logo here when available

OUTPUT_PATH = "/opt/iatc-thin-client/static/description/IA_Treasury_Control_User_Guide.docx"
GUIDE_PATH  = "/opt/iatc-thin-client/static/description/USER_GUIDE.md"

# ── Helpers ───────────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    """Convert RGBColor (a bytes tuple) to hex string like '1B3A6B'."""
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    """Set cell background shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    tcPr.append(shd)


def set_para_bg(para, rgb: RGBColor):
    """Set paragraph shading (for code / blockquote blocks)."""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    pPr.append(shd)


def add_border_para(para, border_color="1B3A6B", border_size=4, side='bottom'):
    """Add a border to a paragraph (used for H1 underline)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bd   = OxmlElement(f'w:{side}')
    bd.set(qn('w:val'),   'single')
    bd.set(qn('w:sz'),    str(border_size))
    bd.set(qn('w:space'), '4')
    bd.set(qn('w:color'), border_color)
    pBdr.append(bd)
    pPr.append(pBdr)


def add_cell_border(cell, sides=('top','bottom','left','right'), color="CCCCCC", size=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tblBd = OxmlElement('w:tcBorders')
    for side in sides:
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'),   'single')
        bd.set(qn('w:sz'),    str(size))
        bd.set(qn('w:space'), '0')
        bd.set(qn('w:color'), color)
        tblBd.append(bd)
    tcPr.append(tblBd)


def set_table_borders(table, color="CCCCCC"):
    for row in table.rows:
        for cell in row.cells:
            add_cell_border(cell, color=color)


def para_space(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)


def apply_inline_formatting(run_text, para):
    """Parse **bold**, `code` inline and add as runs to para."""
    # Split on **bold** and `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts   = pattern.split(run_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(10)
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            r = para.add_run(part)
            r.font.size = Pt(10)
    return para


def page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)


# ── Header / Footer ───────────────────────────────────────────────────────────

def setup_header_footer(doc):
    """Add header (logos) and footer (copyright) to the default section."""
    section = doc.sections[0]
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    # ── Header ────────────────────────────────────────────────────────────────
    header = section.header
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()

    # Create a 1-row, 2-col table for left/right logos
    htbl = header.add_table(1, 2, width=Cm(17))
    htbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove table borders
    tblPr = OxmlElement('w:tblPr')
    tblBrd = OxmlElement('w:tblBorders')
    for side in ('top','bottom','left','right','insideH','insideV'):
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'none')
        bd.set(qn('w:sz'), '0')
        bd.set(qn('w:space'), '0')
        bd.set(qn('w:color'), 'auto')
        tblBrd.append(bd)
    tblPr.append(tblBrd)
    htbl._tbl.insert(0, tblPr)

    left_cell  = htbl.rows[0].cells[0]
    right_cell = htbl.rows[0].cells[1]

    # Left cell: Uniasser logo
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run()
    lr.add_picture(LOGO_PATH, width=Cm(2.5))

    # Right cell: Vantis placeholder
    # TODO: When Vantis logo file is available, replace the text below with:
    # rr.add_picture(VANTIS_LOGO_PATH, width=Cm(2.5))
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run("[ Vantis logo ]")
    rr.font.size  = Pt(9)
    rr.font.color.rgb = GRAY_TEXT
    rr.font.italic    = True

    # ── Footer ────────────────────────────────────────────────────────────────
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("IA Treasury Control v1.7.0  ·  © 2026 Uniasser  ·  OPL-1 License")
    fr.font.size  = Pt(8)
    fr.font.color.rgb = GRAY_TEXT


# ── Title page ────────────────────────────────────────────────────────────────

def add_title_page(doc):
    # Big spacer
    for _ in range(6):
        sp = doc.add_paragraph()
        para_space(sp, 0, 0)

    # Main title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(tp, 0, 6)
    tr = tp.add_run("IA Treasury Control")
    tr.bold           = True
    tr.font.size      = Pt(32)
    tr.font.color.rgb = NAVY

    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(tp2, 0, 4)
    tr2 = tp2.add_run("MCP — Complete User Guide")
    tr2.bold           = True
    tr2.font.size      = Pt(22)
    tr2.font.color.rgb = BLUE

    # Divider line
    dl = doc.add_paragraph()
    dl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_para(dl, side='bottom')
    para_space(dl, 12, 12)

    # Version / compatibility
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(vp, 0, 4)
    vr = vp.add_run("Version 1.7.0")
    vr.font.size      = Pt(13)
    vr.font.color.rgb = GRAY_TEXT

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(cp, 0, 4)
    cr = cp.add_run("Compatibility: Odoo 16, 17, 18, 19 and Odoo.sh")
    cr.font.size      = Pt(11)
    cr.font.color.rgb = GRAY_TEXT

    yp = doc.add_paragraph()
    yp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(yp, 0, 4)
    yr = yp.add_run("© 2026 Uniasser · OPL-1 License")
    yr.font.size      = Pt(10)
    yr.font.color.rgb = GRAY_TEXT
    yr.font.italic    = True

    # Page break after title
    doc.add_page_break()


# ── Table of Contents page ────────────────────────────────────────────────────

TOC_ENTRIES = [
    ("1",   "Prerequisites"),
    ("2",   "Get your license"),
    ("3",   "Install the module in Odoo"),
    ("4",   "Initial configuration"),
    ("5",   "Connect with Claude (claude.ai)"),
    ("6",   "Connect with Gemini (Google)"),
    ("7",   "Connect with ChatGPT (OpenAI)"),
    ("8",   "Connect with WhatsApp (Twilio)"),
    ("8b",  "Free Trial & Your Own Anthropic API Key"),
    ("9",   "WhatsApp vs. AI Chat — Capabilities"),
    ("10",  "Available tools"),
    ("11",  "Usage examples"),
    ("12",  "Auto-install scripts"),
    ("13",  "Troubleshooting"),
    ("14",  "Ubuntu Linux 24.04 — Self-hosted notes"),
]

def add_toc_page(doc):
    h = doc.add_paragraph()
    hr = h.add_run("Table of Contents")
    hr.bold           = True
    hr.font.size      = Pt(18)
    hr.font.color.rgb = NAVY
    para_space(h, 0, 12)
    add_border_para(h, side='bottom')

    for num, title in TOC_ENTRIES:
        tp = doc.add_paragraph()
        para_space(tp, 2, 2)
        tp.paragraph_format.left_indent = Cm(0.3)
        nr = tp.add_run(f"{str(num)}.  ")
        nr.font.size      = Pt(10)
        nr.font.color.rgb = NAVY
        nr.bold           = True
        tr = tp.add_run(title)
        tr.font.size      = Pt(10)

    doc.add_page_break()


# ── Markdown parser + renderer ────────────────────────────────────────────────

def strip_md_links(text):
    """Replace [label](url) with label."""
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def add_h1(doc, text):
    text = strip_md_links(text)
    p = doc.add_paragraph()
    page_break_before(p)
    para_space(p, 0, 4)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(16)
    r.font.color.rgb = NAVY
    add_border_para(p, side='bottom', border_size=6)
    return p


def add_h2(doc, text):
    text = strip_md_links(text)
    p = doc.add_paragraph()
    para_space(p, 10, 4)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(13)
    r.font.color.rgb = BLUE
    return p


def add_h3(doc, text):
    text = strip_md_links(text)
    p = doc.add_paragraph()
    para_space(p, 8, 3)
    r = p.add_run(text)
    r.bold           = True
    r.font.size      = Pt(11)
    r.font.color.rgb = DARKGRAY
    return p


def add_body(doc, text, indent=0):
    text = strip_md_links(text)
    p = doc.add_paragraph()
    para_space(p, 1, 1)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    apply_inline_formatting(text, p)
    return p


def add_bullet(doc, text, level=0):
    text = strip_md_links(text)
    p = doc.add_paragraph(style='List Bullet')
    para_space(p, 1, 1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    apply_inline_formatting(text, p)
    return p


def add_blockquote(doc, text):
    text = strip_md_links(text)
    # Remove leading > and emoji shortcodes
    text = re.sub(r'^>\s*', '', text)
    p = doc.add_paragraph()
    para_space(p, 3, 3)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    set_para_bg(p, QUOTE_BG)
    apply_inline_formatting(text, p)
    for run in p.runs:
        run.font.italic = True
        run.font.size   = Pt(10)
    return p


def add_code_block(doc, lines, lang=''):
    for line in lines:
        p = doc.add_paragraph()
        para_space(p, 0, 0)
        p.paragraph_format.left_indent = Cm(0.5)
        set_para_bg(p, CODE_BG)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        # thin border on all sides for first/last lines would be complex — skip
    # Add a tiny space after the block
    sp = doc.add_paragraph()
    para_space(sp, 0, 4)
    return


def add_md_table(doc, header_row, rows):
    """Render a markdown table as a Word table."""
    cols = len(header_row)
    tbl  = doc.add_table(rows=1, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = 'Table Grid'

    # Header row
    hdr = tbl.rows[0]
    for i, cell_text in enumerate(header_row):
        cell = hdr.cells[i]
        set_cell_bg(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(strip_md_links(cell_text.strip()))
        r.bold           = True
        r.font.size      = Pt(9)
        r.font.color.rgb = WHITE

    # Data rows
    for idx, row_data in enumerate(rows):
        row = tbl.add_row()
        bg  = LIGHT_BLUE_BG if idx % 2 == 1 else WHITE
        for i, cell_text in enumerate(row_data):
            if i >= cols:
                break
            cell = row.cells[i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_inline_formatting(strip_md_links(cell_text.strip()), p)
            for run in p.runs:
                run.font.size = Pt(9)

    set_table_borders(tbl)
    sp = doc.add_paragraph()
    para_space(sp, 0, 6)


def parse_table_row(line):
    """Split a markdown table row into cells."""
    line = line.strip().strip('|')
    return [c.strip() for c in line.split('|')]


def is_separator_row(row):
    return all(re.match(r'^:?-+:?$', c.strip()) for c in row if c.strip())


# ── Main markdown → docx ──────────────────────────────────────────────────────

def render_markdown(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    n = len(lines)

    # Buffer for table rows
    table_buf = []   # list of lists of cells
    in_table  = False

    # Buffer for code block
    in_code   = False
    code_buf  = []
    code_lang = ''

    # Buffer for blockquote
    in_quote  = False
    quote_buf = []

    def flush_table():
        nonlocal table_buf, in_table
        if len(table_buf) >= 2:
            header = table_buf[0]
            # find separator
            data_start = 1
            if len(table_buf) > 1 and is_separator_row(table_buf[1]):
                data_start = 2
            add_md_table(doc, header, table_buf[data_start:])
        table_buf = []
        in_table  = False

    def flush_quote():
        nonlocal quote_buf, in_quote
        for ql in quote_buf:
            add_blockquote(doc, ql)
        quote_buf = []
        in_quote  = False

    while i < n:
        raw  = lines[i]
        line = raw.rstrip()

        # ── Code block ──────────────────────────────────────────────────────
        if line.startswith('```'):
            if in_table:  flush_table()
            if in_quote:  flush_quote()
            if not in_code:
                in_code   = True
                code_lang = line[3:].strip()
                code_buf  = []
            else:
                add_code_block(doc, code_buf, code_lang)
                in_code   = False
                code_buf  = []
                code_lang = ''
            i += 1
            continue

        if in_code:
            code_buf.append(raw.rstrip('\n'))
            i += 1
            continue

        # ── Table ────────────────────────────────────────────────────────────
        if line.startswith('|'):
            if in_quote: flush_quote()
            cells = parse_table_row(line)
            table_buf.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Blockquote ───────────────────────────────────────────────────────
        if line.startswith('>'):
            if in_table: flush_table()
            quote_buf.append(line)
            in_quote = True
            i += 1
            continue
        else:
            if in_quote:
                flush_quote()

        # ── Headings ─────────────────────────────────────────────────────────
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            if level == 1:
                add_h1(doc, text)
            elif level == 2:
                add_h2(doc, text)
            else:
                add_h3(doc, text)
            i += 1
            continue

        # ── HR ───────────────────────────────────────────────────────────────
        if re.match(r'^-{3,}$', line) or re.match(r'^\*{3,}$', line):
            i += 1
            continue  # skip decorative HRs

        # ── Bullet lists ─────────────────────────────────────────────────────
        mb = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if mb:
            level = len(mb.group(1)) // 2
            add_bullet(doc, mb.group(2), level=level)
            i += 1
            continue

        # ── Numbered lists ────────────────────────────────────────────────────
        mn = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if mn:
            indent = len(mn.group(1)) // 3
            add_bullet(doc, mn.group(2), level=indent)
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Italic-only line (e.g. footer *...*) ─────────────────────────────
        if re.match(r'^\*[^*].*[^*]\*$', line):
            p = doc.add_paragraph()
            para_space(p, 2, 2)
            r = p.add_run(line.strip('*'))
            r.italic    = True
            r.font.size = Pt(9)
            r.font.color.rgb = GRAY_TEXT
            i += 1
            continue

        # ── Plain body text ───────────────────────────────────────────────────
        add_body(doc, line)
        i += 1

    # Flush any remaining buffers
    if in_table: flush_table()
    if in_quote: flush_quote()
    if in_code:  add_code_block(doc, code_buf, code_lang)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Read the markdown file
    guide = Path(GUIDE_PATH).read_text(encoding='utf-8')

    # Extract the English section
    marker = "# IA Treasury Control (MCP) — Complete User Guide"
    idx = guide.find(marker)
    if idx == -1:
        raise ValueError("English section marker not found in USER_GUIDE.md")
    english_md = guide[idx:]

    # Create document
    doc = Document()

    # Page setup: 2 cm margins all sides
    for section in doc.sections:
        section.top_margin    = Cm(2.5)   # extra for header
        section.bottom_margin = Cm(2.5)   # extra for footer
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # Header + footer on every page
    setup_header_footer(doc)

    # Title page
    add_title_page(doc)

    # Table of Contents
    add_toc_page(doc)

    # Render the markdown (skip the H1 title line — already on title page)
    lines  = english_md.split('\n')
    # Drop first line (the H1 title) and the next few metadata lines
    # We keep the body from "## Table of Contents" onward
    # Actually let's keep everything so the TOC header is rendered
    render_markdown(doc, english_md)

    # Save
    out = Path(OUTPUT_PATH)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"✅ Document saved: {out}")
    print(f"   Size: {out.stat().st_size:,} bytes")


if __name__ == '__main__':
    main()
